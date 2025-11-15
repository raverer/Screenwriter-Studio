# app.py — Screenwriter Studio — Groq-backed (Streamlit Cloud ready)
# Full app adapted from your original file to use Groq API for generation.

import os
import re
import json
import html
from datetime import datetime
from typing import Any, List, Dict, Optional, Iterable

import streamlit as st

# -------------------- Groq client (cloud) --------------------
try:
    from groq import Groq
except Exception:
    Groq = None

# Optional libraries (docx, reportlab)
try:
    from docx import Document
    from docx.shared import Inches
except Exception:
    Document = None
    Inches = None

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas as pdf_canvas
except Exception:
    letter = None
    pdf_canvas = None

# -------------------- Page config --------------------
st.set_page_config(page_title="Screenwriter Studio", layout="wide")

# -------------------- Session defaults --------------------
st.session_state.setdefault("sidebar_collapsed", False)
st.session_state.setdefault("script_text", "")
st.session_state.setdefault("chat_history", [])
st.session_state.setdefault("last_assistant", "")
st.session_state.setdefault("story_outline", "")
st.session_state.setdefault("writing_style", st.session_state.get("writing_style", ""))
st.session_state.setdefault("trend_style", st.session_state.get("trend_style", ""))
st.session_state.setdefault("current_project", st.session_state.get("current_project", ""))
st.session_state.setdefault("clear_chat_box", False)
st.session_state.setdefault("force_rerun", False)

# -------------------- Simple theme CSS --------------------
st.markdown(
    """
<style>
:root{ --bg:#ffffff; --text:#111111; --panel:#f4f5f7; --border:#e0e4ea; --card:#fbfbfb; --muted:#6b7280; --accent:#1e3a8a; --accent-hover:#1e40af; }
@media (prefers-color-scheme: dark){ :root{ --bg:#071019; --text:#e6eef6; --panel:#0f1724; --border:#21303f; --card:#0b1220; --muted:#93a1b1; --accent:#3b82f6; --accent-hover:#60a5fa; } }
html, body, .main, .block-container { background: var(--bg) !important; color: var(--text) !important; }
[data-testid="stSidebar"] { background: var(--panel) !important; color: var(--text) !important; width: 320px !important; padding: 16px !important; border-right: 1px solid var(--border) !important; }
.main-content { margin-left: 320px !important; padding: 20px; }
.screenplay-box { background: var(--card) !important; color: var(--text) !important; min-height: 72vh; padding: 20px; border-radius: 8px; border: 1px solid var(--border); white-space: pre-wrap; font-family: "Courier New", monospace; font-size:14px; line-height:1.45; }
.chat-box { background: var(--card) !important; color: var(--text) !important; min-height: 72vh; padding: 16px; border-radius: 8px; border: 1px solid var(--border); font-size: 15px; }
button, .stButton>button { background: var(--accent) !important; color: white !important; border-radius: 6px !important; border: none !important; padding: 8px 14px !important; font-weight:600 !important; }
button:hover, .stButton>button:hover { background: var(--accent-hover) !important; }
input, textarea, select { background: transparent !important; color: var(--text) !important; border: 1px solid var(--border) !important; border-radius:6px !important; padding:6px !important; }
.small-muted { color: var(--muted) !important; font-size:13px; }
</style>
""",
    unsafe_allow_html=True,
)

# -------------------- Sidebar toggle --------------------
toggle_container = st.container()
with toggle_container:
    c1, c2 = st.columns([0.04, 0.96])
    with c1:
        label = "»" if st.session_state["sidebar_collapsed"] else "«"
        if st.button(label, key="sidebar_toggle"):
            st.session_state["sidebar_collapsed"] = not st.session_state["sidebar_collapsed"]
            st.session_state["force_rerun"] = True

if st.session_state["sidebar_collapsed"]:
    st.markdown("""<style>[data-testid="stSidebar"]{transform:translateX(-360px)!important;opacity:0}</style>""", unsafe_allow_html=True)
else:
    st.markdown("""<style>[data-testid="stSidebar"]{transform:translateX(0)!important;opacity:1}</style>""", unsafe_allow_html=True)

st.markdown('<div class="main-content">', unsafe_allow_html=True)

# -------------------- Sidebar controls --------------------
st.sidebar.header("⚙️ Project Controls")
proj_col1, proj_col2 = st.sidebar.columns([2, 1])
with proj_col1:
    project_name = st.text_input("Project name", value=st.session_state.get("current_project", ""), key="project_name_input")
with proj_col2:
    if st.button("New", key="proj_new_btn"):
        new_name = f"project_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        st.session_state["current_project"] = new_name
        st.session_state["script_text"] = ""
        st.session_state["chat_history"] = []
        st.session_state["story_outline"] = ""
        st.session_state["force_rerun"] = True

proj_dir = "projects"
os.makedirs(proj_dir, exist_ok=True)
existing_projects = [f[:-5] for f in os.listdir(proj_dir) if f.endswith(".json")]
selected_project = st.sidebar.selectbox("Load project", ["(none)"] + sorted(existing_projects), key="proj_load_select")
if selected_project and selected_project != "(none)":
    if st.sidebar.button("Load", key="proj_load_btn"):
        path = os.path.join(proj_dir, selected_project + ".json")
        try:
            with open(path, "r", encoding="utf-8") as fh:
                data = json.load(fh)
            st.session_state["script_text"] = data.get("script_text", "")
            st.session_state["chat_history"] = data.get("chat_history", [])
            st.session_state["current_project"] = selected_project
            st.session_state["story_outline"] = data.get("story_outline", "")
            st.success(f"Loaded {selected_project}")
            st.session_state["force_rerun"] = True
        except Exception as e:
            st.error(f"Load failed: {e}")

if st.sidebar.button("Save project", key="proj_save_btn"):
    name = (project_name.strip() or st.session_state.get("current_project") or f"project_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
    path = os.path.join(proj_dir, name + ".json")
    payload = {
        "script_text": st.session_state.get("script_text", ""),
        "chat_history": st.session_state.get("chat_history", []),
        "story_outline": st.session_state.get("story_outline", ""),
        "saved_at": datetime.now().isoformat(),
    }
    with open(path, "w", encoding="utf-8") as fh:
        json.dump(payload, fh, ensure_ascii=False, indent=2)
    st.session_state["current_project"] = name
    st.success(f"Saved project: {name}")

st.sidebar.markdown("---")

# -------------------- Generation controls & metadata --------------------
temperature = st.sidebar.slider("Temperature", 0.1, 1.2, 0.8, key="temperature_slider")
ctx_size = st.sidebar.number_input("Context length", 512, 65536, 4096, step=256, key="context_length_input")

st.sidebar.markdown("### Metadata")
title = st.sidebar.text_input("Title / Idea", value=st.session_state.get("meta_title", ""), key="meta_title")
genre = st.sidebar.text_input("Genre", value=st.session_state.get("meta_genre", ""), key="meta_genre")
tone = st.sidebar.text_input("Tone / Style", value=st.session_state.get("meta_tone", ""), key="meta_tone")
characters = st.sidebar.text_input("Characters", value=st.session_state.get("meta_characters", ""), key="meta_characters")
setting = st.sidebar.text_input("Setting", value=st.session_state.get("meta_setting", ""), key="meta_setting")
duration = st.sidebar.text_input("Duration (optional)", key="meta_duration")
notes = st.sidebar.text_area("Extra Notes", height=120, value=st.session_state.get("meta_notes", ""), key="meta_notes")

script_format = st.sidebar.selectbox(
    "Script Format",
    options=["Short Film", "Social Media Viral Short (30–90 sec)"],
    index=0,
    help="Short Film for festival/student shorts; Social Media for viral formats.",
    key="script_format_select",
)
st.session_state["script_format"] = script_format

trend_style = st.sidebar.text_input("Trend / Viral Style (optional)", value=st.session_state.get("trend_style", ""), help="Examples: 'Gen Z humor, fast cuts', 'ASMR aesthetic', 'emotional tearjerker', 'product showcase'", key="trend_style_input")
st.session_state["trend_style"] = trend_style

writing_style = st.sidebar.text_input("Writing Style (optional)", value=st.session_state.get("writing_style", ""), help="Example: Quentin Tarantino, Studio Ghibli, Martin Scorsese, Coen Brothers, Wes Anderson", key="writing_style_input")
st.session_state["writing_style"] = writing_style

st.sidebar.markdown("---")
uploaded_template = st.sidebar.file_uploader("Upload DOCX template (optional)", type=["docx"], key="uploaded_template")
use_template = st.sidebar.checkbox("Use uploaded template", value=False, key="use_template_chk")
gen_btn = st.sidebar.button("Generate Script", key="gen_btn")
gen_next_btn = st.sidebar.button("Generate Next Scene", key="gen_next_btn")
export_docx_btn = st.sidebar.button("Export DOCX", key="export_docx_btn")
export_pdf_btn = st.sidebar.button("Export PDF", key="export_pdf_btn")
st.sidebar.markdown("<div class='small-muted'>Use Save project to persist work.</div>", unsafe_allow_html=True)

# -------------------- SMART SYSTEM PROMPT --------------------
SYSTEM_SMART = """
You are an award-winning screenwriter. Write in clean, professional screenplay format (SCENE HEADINGS, ACTION, CHARACTER, DIALOGUE).

Core rules:
1. Maintain story logic, escalation, and clear character motivations.
2. Match the user's metadata: genre, tone, style, setting, characters, and notes.
3. If a writing style is given, blend it subtly into pacing, dialogue, rhythm, and visuals.
4. If a viral/short format is chosen, keep scenes punchy, high-impact, and focused on a single emotional or comedic payoff.
5. Avoid repetition, filler, or explaining what you are doing.
6. Think internally, but output only the screenplay.
7. Keep dialogue natural, cinematic, and consistent with character voices.
"""

# -------------------- AI SCRIPT DOCTOR --------------------
script_doctor_btn = st.sidebar.button("AI Script Doctor")
doctor_prompt = f"""
You are an expert screenwriting consultant.
Analyze the user's script and return:

1. Fixes for plot holes
2. Stronger character motivations
3. Better pacing
4. Improvements to scene transitions
5. Dialogue polish
6. Suggestions for visual storytelling

Then propose a rewritten improved version of the weak sections.
Here is the script:

{st.session_state["script_text"]}
"""

# -------------------- Utilities --------------------
def safe_str(x: Any) -> str:
    try:
        return str(x)
    except Exception:
        return ""

def extract_from_resp(resp: Any) -> str:
    if isinstance(resp, Iterable) and not isinstance(resp, (str, bytes, dict)):
        try:
            parts = []
            for p in resp:
                parts.append(safe_str(p))
            return "".join(parts)
        except Exception:
            pass
    if hasattr(resp, "message"):
        msg = getattr(resp, "message")
        if isinstance(msg, dict):
            return safe_str(msg.get("content", ""))
        return safe_str(msg)
    if isinstance(resp, dict):
        if "message" in resp and isinstance(resp["message"], dict) and "content" in resp["message"]:
            return safe_str(resp["message"]["content"])
        if "content" in resp:
            return safe_str(resp["content"])
        return safe_str(resp)
    return safe_str(resp)

def clean_model_text(raw: Any) -> str:
    t = extract_from_resp(raw)
    m = re.search(r'content=(["\'])(.*?)\1', t, flags=re.DOTALL)
    if m:
        t = m.group(2)
    t = t.replace("\\r\\n", "\n").replace("\\n", "\n").replace("\\r", "\n")
    t = re.sub(r'<[^>]+>', '', t)
    t = re.sub(r'^\s*Here is the screenplay.*', '', t, flags=re.IGNORECASE)
    t = t.strip()
    t = re.sub(r'\n{3,}', '\n\n', t)
    return t

def safe_html_from_text(text: str) -> str:
    escaped = html.escape(text or "")
    return escaped.replace("\n", "<br>")

# -------------------- Script parsing --------------------
def parse_script_blocks(text: str) -> List[Dict[str, str]]:
    if not text:
        return []
    text = text.replace("\\n", "\n")
    lines = [ln.rstrip() for ln in text.splitlines()]
    blocks = []
    i = 0
    while i < len(lines):
        ln = lines[i].strip()
        if ln == "":
            i += 1
            continue
        if re.match(r'^(INT\.|EXT\.|INT/EXT\.)', ln, flags=re.IGNORECASE):
            blocks.append({"type": "scene_heading", "text": ln})
            i += 1
            continue
        if re.search(r'(CUT TO:|FADE OUT:|FADE IN:)$', ln, flags=re.IGNORECASE):
            blocks.append({"type": "transition", "text": ln})
            i += 1
            continue
        if ln.isupper() and len(ln.split()) <= 6:
            blocks.append({"type": "character", "text": ln})
            j = i + 1
            dlg = []
            while j < len(lines) and lines[j].strip() and not lines[j].strip().isupper():
                dlg.append(lines[j].strip())
                j += 1
            if dlg:
                blocks.append({"type": "dialogue", "text": " ".join(dlg)})
            i = j
            continue
        blocks.append({"type": "action", "text": ln})
        i += 1
    return blocks

# -------------------- Auto model selection (unchanged) --------------------
def pick_model_for_generation(prompt_text: str) -> str:
    if len(prompt_text or "") < 300:
        return "llama3-8b-8192"
    return "llama3-70b-8192"

# -------------------- Groq wrapper --------------------
def call_groq_safe(system_msg: str, user_msg: str, model_name: str, stream: bool=False, add_to_chat: bool=True) -> str:
    """
    Use Groq cloud API via groq-python client.
    - Reads GROQ_API_KEY from st.secrets or environment.
    - Sends system + user as a single user prompt (light context injection).
    - Does NOT re-feed full chat history, only keeps it locally for UI.
    """
    if Groq is None:
        err = "Groq client library not installed. Add 'groq' to requirements."
        if add_to_chat:
            st.session_state["chat_history"].append({"role":"assistant","content":err})
            st.session_state["last_assistant"] = err
        return err

    # read key (Streamlit secrets preferred)
    groq_key = None
    try:
        groq_key = st.secrets["GROQ_API_KEY"]
    except Exception:
        groq_key = os.environ.get("GROQ_API_KEY")

    if not groq_key:
        err = "GROQ_API_KEY not found. Set it in Streamlit secrets or environment."
        if add_to_chat:
            st.session_state["chat_history"].append({"role":"assistant","content":err})
            st.session_state["last_assistant"] = err
        return err

    # Initialize client
    try:
        client = Groq(api_key=groq_key)
    except Exception:
        try:
            client = Groq()
        except Exception as e:
            err = f"Failed to init Groq client: {e}"
            if add_to_chat:
                st.session_state["chat_history"].append({"role":"assistant","content":err})
                st.session_state["last_assistant"] = err
            return err

    # Minimal context injection (last n chars of script)
    context = st.session_state.get("script_text", "")[-4000:]

    # Build single prompt message (system + user + context)
    # We keep system_msg as guidance and include context for reference
    user_payload = (
        f"{system_msg}\n\n"
        f"CONTEXT (use only if helpful, do not repeat):\n{context}\n\n"
        f"USER REQUEST:\n{user_msg}"
    )

    try:
        # Use the chat completion API (OpenAI-compatible endpoint)
        response = client.chat.completions.create(
            model=model_name,
            messages=[{"role":"user","content": user_payload}],
            temperature=float(temperature),
            max_tokens=8192  # allow long generation; Groq enforces model limits
        )
        # Extract text
        content = ""
        try:
            content = response.choices[0].message.content
        except Exception:
            content = str(response)
        raw = clean_model_text(content)

    except Exception as e:
        raw = f"Error calling Groq: {e}"

    cleaned = raw

    # Record minimal chat history locally (not re-fed to model)
    if add_to_chat:
        st.session_state["chat_history"].append({"role":"user","content": user_msg})
        st.session_state["chat_history"].append({"role":"assistant","content": cleaned})

    st.session_state["last_assistant"] = cleaned
    return cleaned

# For backward compatibility in the app, alias call_ollama_safe -> call_groq_safe
call_ollama_safe = call_groq_safe

# -------------------- Prompt builders (unchanged) --------------------
def build_generation_prompt(metadata: Dict[str, str]) -> str:
    lines = [
        f"Title: {metadata.get('title','')}",
        f"Genre: {metadata.get('genre','')}",
        f"Tone: {metadata.get('tone','')}",
        f"Format: {metadata.get('format','')}",
        f"Duration: {metadata.get('duration','')}",
        f"Characters: {metadata.get('characters','')}",
        f"Setting: {metadata.get('setting','')}",
        f"Notes: {metadata.get('notes','')}",
        "",
        "Write a screenplay in standard format. Use INT./EXT. headings, action, CHARACTER names in caps, and dialogue. Keep it cinematic.",
    ]
    return "\n".join(lines)

# --- Viral & shortfilm prompt functions (same as your existing ones) ---
def build_viral_outline_prompt(meta: Dict[str, str]) -> str:
    return f"""
{SYSTEM_SMART}

Create a concise 4-beat micro-outline optimized for short-form social video (30-90 seconds).
METADATA:
• Title/Idea: {meta.get('title','')}
• Genre: {meta.get('genre','')}
• Tone: {meta.get('tone','')}
• Characters: {meta.get('characters','')}
• Setting: {meta.get('setting','')}
• Notes: {meta.get('notes','')}
• Viral style: {meta.get('trend','None')}
• Optional writing style: {meta.get('style','None')}

REQUIREMENTS:
1) Output a 4-beat outline: HOOK (0-3s), BUILD (3-25s), PUNCHLINE/CLIMAX (25-50s), LOOP/CTA (final few seconds).
2) Keep it highly visual with a thumb-stopping first 3 seconds.
3) Provide 2 short variations (A/B) for the hook.
4) Output ONLY the outline (labeled).
"""

def build_viral_script_prompt(meta: Dict[str, str], outline: str) -> str:
    return f"""
{SYSTEM_SMART}

Using the outline below, write a short-form social video script (30-90s) optimized for platform retention and virality.

OUTLINE:
{outline}

RULES:
1) Script must be 3-6 beats, roughly 30-90 seconds.
2) Provide: (a) Final script (camera directions inline), (b) Shot ideas (bullet list), (c) Suggested caption (1-2 lines), (d) 6-10 hashtags.
3) Include two variations: Variation A (safer), Variation B (edgier).
4) Include a loopable ending or micro-CTA.
5) Keep language punchy and highly visual.
6) Output ONLY the requested sections with clear labels.
"""

def build_shortfilm_outline_prompt(meta: Dict[str, str]) -> str:
    return f"""
{SYSTEM_SMART}

Produce a 6-beat story outline for a short film.

METADATA:
• Title/Idea: {meta.get('title','')}
• Genre: {meta.get('genre','')}
• Tone: {meta.get('tone','')}
• Characters: {meta.get('characters','')}
• Setting: {meta.get('setting','')}
• Notes: {meta.get('notes','')}
• Optional writing style: {meta.get('style','None')}

REQUIREMENTS:
1) Provide 6 beats: Opening Image, Setup, Inciting Incident, Rising Conflict, Climax, Resolution.
2) Each beat: 1-2 sentences describing action and emotional stakes.
3) Include protagonist's emotional arc.
4) Output ONLY the outline.
"""

def build_shortfilm_script_prompt(meta: Dict[str, str], outline: str) -> str:
    return f"""
{SYSTEM_SMART}

Using the outline below, write a COMPLETE short-film screenplay from FADE IN to FADE OUT.

OUTLINE:
{outline}

REQUIREMENTS:
1) Write a FULL screenplay. Do NOT stop early.
2) Each scene must have:
   • Proper scene heading (INT/EXT + LOCATION + TIME)
   • 3–8 lines of action
   • Character-consistent dialogue
   • Emotional progression and rising tension
3) The screenplay MUST end with a proper final scene and a clear resolution.
4) Finish with: FADE OUT.
5) Output ONLY the screenplay.
"""

# -------------------- Generators & orchestration (same logic) --------------------
def generate_outline_shortfilm():
    meta = {
        "title": st.session_state.get("meta_title", ""),
        "genre": st.session_state.get("meta_genre", ""),
        "tone": st.session_state.get("meta_tone", ""),
        "characters": st.session_state.get("meta_characters", ""),
        "setting": st.session_state.get("meta_setting", ""),
        "notes": st.session_state.get("meta_notes", ""),
        "style": st.session_state.get("writing_style", ""),
    }
    prompt = build_shortfilm_outline_prompt(meta)
    model = "llama3-8b-8192"
    with st.spinner("Generating story outline (Short Film)…"):
        result = call_groq_safe(SYSTEM_SMART, prompt, model, stream=False, add_to_chat=False)
    st.session_state["story_outline"] = result
    return result

def generate_outline_viral():
    meta = {
        "title": st.session_state.get("meta_title", ""),
        "genre": st.session_state.get("meta_genre", ""),
        "tone": st.session_state.get("meta_tone", ""),
        "characters": st.session_state.get("meta_characters", ""),
        "setting": st.session_state.get("meta_setting", ""),
        "notes": st.session_state.get("meta_notes", ""),
        "trend": st.session_state.get("trend_style", ""),
        "style": st.session_state.get("writing_style", ""),
    }
    prompt = build_viral_outline_prompt(meta)
    model = "llama3-8b-8192"
    with st.spinner("Generating viral outline…"):
        result = call_groq_safe(SYSTEM_SMART, prompt, model, stream=False, add_to_chat=False)
    st.session_state["story_outline"] = result
    return result

def generate_script_shortfilm():
    outline = st.session_state.get("story_outline", "").strip()
    if not outline:
        outline = generate_outline_shortfilm()
    meta = {
        "title": st.session_state.get("meta_title", ""),
        "genre": st.session_state.get("meta_genre", ""),
        "tone": st.session_state.get("meta_tone", ""),
        "characters": st.session_state.get("meta_characters", ""),
        "setting": st.session_state.get("meta_setting", ""),
        "duration": st.session_state.get("meta_duration", ""),
        "notes": st.session_state.get("meta_notes", ""),
        "style": st.session_state.get("writing_style", ""),
    }
    prompt = build_shortfilm_script_prompt(meta, outline)
    model = "llama3-70b-8192"
    with st.spinner("Generating short film screenplay…"):
        result = call_groq_safe(SYSTEM_SMART, prompt, model, stream=False, add_to_chat=False)
    st.session_state["script_text"] = result
    st.session_state["last_assistant"] = result
    return result

def generate_script_viral():
    outline = st.session_state.get("story_outline", "").strip()
    if not outline:
        outline = generate_outline_viral()
    meta = {
        "title": st.session_state.get("meta_title", ""),
        "genre": st.session_state.get("meta_genre", ""),
        "tone": st.session_state.get("meta_tone", ""),
        "characters": st.session_state.get("meta_characters", ""),
        "setting": st.session_state.get("meta_setting", ""),
        "duration": st.session_state.get("meta_duration", ""),
        "notes": st.session_state.get("meta_notes", ""),
        "style": st.session_state.get("writing_style", ""),
        "trend": st.session_state.get("trend_style", ""),
    }
    prompt = build_viral_script_prompt(meta, outline)
    model = "llama3-70b-8192"
    with st.spinner("Generating viral script…"):
        result = call_groq_safe(SYSTEM_SMART, prompt, model, stream=False, add_to_chat=False)
    st.session_state["script_text"] = result
    st.session_state["last_assistant"] = result
    return result

def generate_full_script():
    fmt = st.session_state.get("script_format", "Short Film")
    if "Short Film" in fmt:
        return generate_script_shortfilm()
    else:
        return generate_script_viral()

def generate_next_scene():
    draft = st.session_state.get("script_text", "").strip()
    blocks = parse_script_blocks(draft)
    last_heading = None
    for blk in reversed(blocks):
        if blk["type"] == "scene_heading":
            last_heading = blk["text"]
            break
    meta = {
        "title": st.session_state.get("meta_title", ""),
        "genre": st.session_state.get("meta_genre", ""),
        "tone": st.session_state.get("meta_tone", ""),
        "characters": st.session_state.get("meta_characters", ""),
        "setting": st.session_state.get("meta_setting", ""),
        "duration": st.session_state.get("meta_duration", ""),
        "notes": st.session_state.get("meta_notes", ""),
        "style": st.session_state.get("writing_style", ""),
    }
    if "Social Media" in st.session_state.get("script_format", ""):
        prompt = f"""
{SYSTEM_SMART}

We are creating short-form viral content. Use the stored outline:
{st.session_state.get("story_outline","")}

The last generated content ended with: {last_heading or 'N/A'}

Task:
Write a new variation/next beat for the short-form video that follows from existing script.
Include: (1) beat text (short lines), (2) shot idea bullets, (3) suggested caption (1 line), (4) 6 hashtags.

Constraints:
- Keep it 10-30 seconds.
- Maintain the same trend style: {meta.get('trend','None')}
- Use writing style: {meta.get('style','None')}
- Output only the requested sections.
"""
        model = pick_model_for_generation(prompt)
        with st.spinner("Generating next short-form beat/variation…"):
            out = call_groq_safe(SYSTEM_SMART, prompt, model, stream=False, add_to_chat=False)
        current = st.session_state.get("script_text", "").strip()
        appended = (current + "\n\n" + out) if current else out
        st.session_state["script_text"] = appended
        st.session_state["last_assistant"] = out
        return
    next_scene_instruction = "Continue with the next scene that follows logically."
    if st.session_state.get("story_outline"):
        scene_count = sum(1 for b in blocks if b["type"] == "scene_heading")
        beat_index = scene_count + 1
        next_scene_instruction = f"Using the previously created outline, write the scene corresponding to beat #{beat_index}."
    prompt_extra = f"The last scene heading was: {last_heading}\n" if last_heading else ""
    base = build_generation_prompt(meta)
    prompt = f"""
{SYSTEM_SMART}

Base Metadata:
{base}

Optional writing style: {st.session_state.get("writing_style","None")}

{prompt_extra}
{next_scene_instruction}

Requirements:
- Continue in professional screenplay format (INT./EXT., ACTION, CHARACTER, DIALOGUE).
- Match tone and genre.
- Ensure escalation and avoid repeating scenes.
- Output only the next scene text.
"""
    model = pick_model_for_generation(prompt)
    with st.spinner("Generating next scene…"):
        out = call_groq_safe(SYSTEM_SMART, prompt, model, stream=False, add_to_chat=False)
    current = st.session_state.get("script_text", "").strip()
    appended = (current + "\n\n" + out) if current else out
    st.session_state["script_text"] = appended
    st.session_state["last_assistant"] = out
    return out

# -------------------- Export helpers & top controls --------------------
def export_docx_from_blocks(blocks: List[Dict[str, str]], path: str, template_path: Optional[str] = None) -> str:
    if Document is None:
        raise RuntimeError("python-docx not installed.")
    doc = Document(template_path) if (template_path and os.path.exists(template_path)) else Document()
    for blk in blocks:
        t = blk["type"]
        txt = blk["text"]
        if t == "scene_heading":
            p = doc.add_paragraph(); p.add_run(txt).bold = True
        elif t == "character":
            p = doc.add_paragraph()
            try:
                p.alignment = 1
            except Exception:
                pass
            p.add_run(txt).bold = True
        elif t == "dialogue":
            p = doc.add_paragraph()
            try:
                p.paragraph_format.left_indent = Inches(1.0)
            except Exception:
                pass
            p.add_run(txt)
        elif t == "transition":
            p = doc.add_paragraph(); p.add_run(txt).bold = True
        else:
            doc.add_paragraph(txt)
    os.makedirs("exports", exist_ok=True)
    doc.save(path)
    return path

def export_pdf_simple(text: str, path: str) -> str:
    if pdf_canvas is None:
        raise RuntimeError("reportlab not installed.")
    c = pdf_canvas.Canvas(path, pagesize=letter)
    width, height = letter
    margin = 40
    y = height - margin
    max_chars = 92
    for line in text.splitlines():
        if not line:
            y -= 12
        else:
            while len(line) > max_chars:
                seg = line[:max_chars]
                c.drawString(margin, y, seg)
                line = line[max_chars:]
                y -= 12
                if y < margin:
                    c.showPage(); y = height - margin
            c.drawString(margin, y, line)
            y -= 12
        if y < margin:
            c.showPage(); y = height - margin
    c.save()
    return path

col1, col2, col3, col4 = st.columns([3, 1, 1, 1])
with col1:
    st.title("📄 Screenwriter Studio")
with col2:
    if st.button("Regenerate Last", key="regen_last_btn"):
        last = st.session_state.get("last_assistant", "")
        if not last:
            st.warning("No assistant reply to regenerate from.")
        else:
            model = pick_model_for_generation(last)
            out = call_groq_safe(SYSTEM_SMART, last, model, stream=False, add_to_chat=False)
            st.session_state["last_assistant"] = out
            if st.session_state.get("chat_history") and st.session_state["chat_history"][-1]["role"] == "assistant":
                st.session_state["chat_history"][-1]["content"] = out
            else:
                st.session_state["chat_history"].append({"role": "assistant", "content": out})
            st.session_state["force_rerun"] = True
with col3:
    if st.button("Apply Last to Draft", key="apply_last_btn"):
        last = st.session_state.get("last_assistant", "")
        if not last:
            st.warning("No assistant reply to apply.")
        else:
            header = "# NEW ADDITION — Assistant\n\n"
            candidate = last.strip()
            cur = st.session_state.get("script_text", "").strip()
            appended = (cur + "\n\n" + header + candidate) if cur else (header + candidate)
            st.session_state["script_text"] = appended
            st.success("Applied assistant reply to draft.")
            st.session_state["force_rerun"] = True
with col4:
    if st.button("Clear Draft", key="clear_draft_btn"):
        st.session_state["script_text"] = ""
        st.success("Cleared master draft.")
        st.session_state["force_rerun"] = True

if gen_btn:
    generate_full_script()
    st.session_state["force_rerun"] = True
if gen_next_btn:
    generate_next_scene()
    st.session_state["force_rerun"] = True

# -------------------- Master Draft & Chat UI --------------------
left_col, right_col = st.columns([2.2, 1], gap="large")
with left_col:
    st.markdown("### 🗂 Master Draft")
    script_text = st.session_state.get("script_text", "")
    safe_html = safe_html_from_text(script_text)
    st.markdown(f"<div class='screenplay-box'>{safe_html}</div>", unsafe_allow_html=True)

with right_col:
    st.markdown("### 💬 Assistant Chat")
    st.markdown("<div class='chat-box'>", unsafe_allow_html=True)
    chat_hist = st.session_state.get("chat_history", [])[-80:]
    if not chat_hist:
        st.markdown("<div class='small-muted'>No conversation yet. Use the box below to ask the assistant to rewrite, continue, or create scenes.</div>", unsafe_allow_html=True)
    else:
        for idx, msg in enumerate(chat_hist):
            role = msg.get("role")
            content = safe_html_from_text(msg.get("content", ""))
            if role == "assistant":
                st.markdown(f"<div style='margin-bottom:12px'><strong>Assistant</strong><br>{content}</div>", unsafe_allow_html=True)
            else:
                st.markdown(f"<div style='margin-bottom:10px'><strong>You</strong><br>{content}</div>", unsafe_allow_html=True)
    st.markdown("---", unsafe_allow_html=True)

    if st.session_state.get("clear_chat_box"):
        st.session_state["chat_input"] = ""
        st.session_state["clear_chat_box"] = False
        st.session_state["force_rerun"] = True

    user_input = st.text_area("Your message to assistant (rewrite, continue, or prompt)", height=140, key="chat_input")

    left_btn, mid_btn, right_btn = st.columns([1,1,1])
    with left_btn:
        if st.button("Send", key="send_msg_btn"):
            if user_input.strip():
                # minimal context injection in call_groq_safe; feed only last script fragment
                model = pick_model_for_generation(user_input)
                out = call_groq_safe(SYSTEM_SMART, user_input, model, stream=False, add_to_chat=True)
                st.session_state["last_assistant"] = out
                st.session_state["clear_chat_box"] = True
                st.session_state["force_rerun"] = True
            else:
                st.warning("Please enter a message.")
    with mid_btn:
        if st.button("Stream (if supported)", key="stream_msg_btn"):
            if user_input.strip():
                model = pick_model_for_generation(user_input)
                out = call_groq_safe(SYSTEM_SMART, user_input, model, stream=False, add_to_chat=True)
                st.session_state["last_assistant"] = out
                st.session_state["clear_chat_box"] = True
                st.session_state["force_rerun"] = True
            else:
                st.warning("Please enter a message.")
    with right_btn:
        if st.button("Apply reply to draft", key="apply_reply_btn"):
            last = st.session_state.get("last_assistant", "")
            if not last:
                st.warning("No assistant reply available.")
            else:
                hdr = "# NEW FROM ASSISTANT\n\n"
                cur = st.session_state.get("script_text", "").strip()
                appended = (cur + "\n\n" + hdr + last) if cur else (hdr + last)
                st.session_state["script_text"] = appended
                st.success("Applied assistant reply to master draft.")
                st.session_state["force_rerun"] = True
    st.markdown("</div>", unsafe_allow_html=True)

# -------------------- Exports --------------------
if export_docx_btn:
    try:
        blocks_for_export = parse_script_blocks(st.session_state.get("script_text", ""))
        os.makedirs("exports", exist_ok=True)
        fmt_tag = "viral" if "Social Media" in st.session_state.get("script_format", "") else "shortfilm"
        safe_title = (title or "script").strip().replace(" ", "_")
        docx_path = f"exports/{safe_title}_{fmt_tag}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        template_path = None
        if uploaded_template and use_template:
            template_path = os.path.join("exports", "template_" + datetime.now().strftime('%Y%m%d_%H%M%S') + ".docx")
            with open(template_path, "wb") as tf:
                tf.write(uploaded_template.getbuffer())
        out = export_docx_from_blocks(blocks_for_export, docx_path, template_path)
        with open(out, "rb") as fh:
            st.download_button("Download DOCX", fh.read(), file_name=os.path.basename(out), key="download_docx_btn")
        st.success(f"Exported DOCX → {out}")
    except Exception as e:
        st.error(f"DOCX export failed: {e}")

if export_pdf_btn:
    try:
        text = st.session_state.get("script_text", "")
        os.makedirs("exports", exist_ok=True)
        fmt_tag = "viral" if "Social Media" in st.session_state.get("script_format", "") else "shortfilm"
        safe_title = (title or "script").strip().replace(" ", "_")
        pdf_path = f"exports/{safe_title}_{fmt_tag}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        out = export_pdf_simple(text, pdf_path)
        with open(out, "rb") as fh:
            st.download_button("Download PDF", fh.read(), file_name=os.path.basename(out), key="download_pdf_btn")
        st.success(f"Exported PDF → {out}")
    except Exception as e:
        st.error(f"PDF export failed: {e}")

# -------------------- Safe global rerun handler --------------------
if st.session_state.get("force_rerun", False):
    st.session_state["force_rerun"] = False
    st.experimental_rerun()

# -------------------- Close main-content wrapper & footer --------------------
st.markdown("</div>", unsafe_allow_html=True)
st.markdown("<div style='opacity:0.7; padding-top:10px; text-align:center;'>Screenwriter Studio</div>", unsafe_allow_html=True)
